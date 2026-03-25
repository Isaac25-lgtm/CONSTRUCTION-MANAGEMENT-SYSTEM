"""Report data assembly and export generation services.

Each report_key maps to a data-assembly function that returns a dict with:
  - title: str
  - headers: list[str]
  - rows: list[list[str|number]]
  - summary: dict (optional extra KPIs)
"""
import io
import csv
from datetime import date
from decimal import Decimal

from django.core.files.base import ContentFile

from apps.projects.models import Project


# ---------------------------------------------------------------------------
# Data assembly helpers (one per report key)
# ---------------------------------------------------------------------------

def _fmt(v):
    """Format a value for export: Decimal->str, None->''."""
    if v is None:
        return ""
    if isinstance(v, Decimal):
        return str(v)
    if isinstance(v, date):
        return v.isoformat()
    return str(v)


def _short_ref(obj) -> str:
    """Return a compact stable reference for models without a human code field."""
    return str(obj.id).split("-")[0].upper()


def assemble_schedule(project: Project) -> dict:
    from apps.scheduling.models import ProjectTask
    tasks = ProjectTask.objects.filter(project=project, is_parent=False).order_by("sort_order")
    headers = ["Code", "Name", "Phase", "Start", "End", "Duration", "Progress %", "Status", "Critical"]
    rows = []
    for t in tasks:
        rows.append([
            t.code, t.name, t.phase or "", _fmt(t.planned_start), _fmt(t.planned_end),
            t.duration_days, t.progress, t.get_status_display(), "Yes" if t.is_critical else "No",
        ])
    return {"title": f"Schedule - {project.name}", "headers": headers, "rows": rows}


def assemble_financial(project: Project) -> dict:
    from apps.cost.models import BudgetLine, Expense
    lines = BudgetLine.objects.filter(project=project).order_by("sort_order")
    headers = ["Code", "Name", "Category", "Budget", "Actual", "Variance"]
    rows = []
    for bl in lines:
        actual = Expense.objects.filter(budget_line=bl).aggregate(t=models.Sum("amount"))["t"] or 0
        rows.append([bl.code, bl.name, bl.get_category_display(), _fmt(bl.budget_amount), _fmt(actual), _fmt(bl.budget_amount - actual)])
    return {"title": f"Financial Summary - {project.name}", "headers": headers, "rows": rows}


def assemble_risk(project: Project) -> dict:
    from apps.risks.models import Risk
    risks = Risk.objects.filter(project=project).order_by("-created_at")
    headers = ["Code", "Title", "Category", "Likelihood", "Impact", "Score", "Status", "Owner"]
    rows = []
    for r in risks:
        rows.append([
            r.code, r.title, r.get_category_display(), r.get_likelihood_display(),
            r.get_impact_display(), r.risk_score, r.get_status_display(), r.owner.get_full_name() if r.owner else "",
        ])
    return {"title": f"Risk Register - {project.name}", "headers": headers, "rows": rows}


def assemble_rfi(project: Project) -> dict:
    from apps.rfis.models import RFI
    rfis = RFI.objects.filter(project=project).order_by("-created_at")
    headers = ["Code", "Subject", "Raised By", "Date Raised", "Due Date", "Status", "Priority"]
    rows = []
    for r in rfis:
        rows.append([
            r.code, r.subject, r.raised_by.get_full_name() if r.raised_by else "",
            _fmt(r.date_raised), _fmt(r.due_date), r.get_status_display(), r.get_priority_display(),
        ])
    return {"title": f"RFI Register - {project.name}", "headers": headers, "rows": rows}


def assemble_changes(project: Project) -> dict:
    from apps.changes.models import ChangeOrder
    changes = ChangeOrder.objects.filter(project=project).order_by("-created_at")
    headers = ["Code", "Title", "Category", "Cost Impact", "Time Impact (days)", "Status"]
    rows = []
    for c in changes:
        rows.append([
            c.code, c.title, c.get_category_display(),
            _fmt(c.cost_impact), c.time_impact_days or 0, c.get_status_display(),
        ])
    return {"title": f"Change Orders - {project.name}", "headers": headers, "rows": rows}


def assemble_punch(project: Project) -> dict:
    from apps.field_ops.models import PunchItem
    items = PunchItem.objects.filter(project=project).order_by("-created_at")
    headers = ["Ref", "Title", "Location", "Priority", "Status", "Assigned To", "Due Date"]
    rows = []
    for p in items:
        rows.append([
            _short_ref(p), p.title, p.location or "", p.get_priority_display(),
            p.get_status_display(), p.assigned_to.get_full_name() if p.assigned_to else "", _fmt(p.due_date),
        ])
    return {"title": f"Punch List - {project.name}", "headers": headers, "rows": rows}


def assemble_procurement(project: Project) -> dict:
    from apps.procurement.models import PurchaseOrder, POItem
    from django.db.models import Sum, F
    pos = PurchaseOrder.objects.filter(project=project).select_related("supplier").order_by("-order_date")
    headers = ["PO Code", "Supplier", "Order Date", "Delivery Date", "Total", "Status"]
    rows = []
    for po in pos:
        total = POItem.objects.filter(purchase_order=po).aggregate(t=Sum(F("quantity") * F("unit_price")))["t"] or 0
        rows.append([po.code, po.supplier.name, _fmt(po.order_date), _fmt(po.delivery_date), _fmt(total), po.get_status_display()])
    return {"title": f"Procurement Summary - {project.name}", "headers": headers, "rows": rows}


def assemble_labour(project: Project) -> dict:
    from apps.labour.models import TimesheetEntry
    entries = TimesheetEntry.objects.filter(project=project).select_related("resource").order_by("-work_date")
    headers = ["Date", "Resource", "Hours", "Overtime", "Description", "Status"]
    rows = []
    for e in entries:
        rows.append([
            _fmt(e.work_date), e.resource.name if e.resource else "", e.hours, e.overtime_hours,
            e.description or "", e.get_status_display(),
        ])
    return {"title": f"Timesheets - {project.name}", "headers": headers, "rows": rows}


def assemble_safety(project: Project) -> dict:
    from apps.field_ops.models import SafetyIncident
    incidents = SafetyIncident.objects.filter(project=project).order_by("-incident_date")
    headers = ["Ref", "Title", "Date", "Type", "Severity", "Status", "Location"]
    rows = []
    for s in incidents:
        rows.append([
            _short_ref(s), s.title, _fmt(s.incident_date), s.get_incident_type_display(),
            s.get_severity_display(), s.get_status_display(), s.location or "",
        ])
    return {"title": f"Safety Report - {project.name}", "headers": headers, "rows": rows}


def assemble_quality(project: Project) -> dict:
    from apps.field_ops.models import QualityCheck
    checks = QualityCheck.objects.filter(project=project).order_by("-check_date")
    headers = ["Ref", "Title", "Date", "Category", "Result", "Location"]
    rows = []
    for q in checks:
        rows.append([
            _short_ref(q), q.title, _fmt(q.check_date), q.get_category_display(),
            q.get_result_display(), q.location or "",
        ])
    return {"title": f"Quality Report - {project.name}", "headers": headers, "rows": rows}


def assemble_meetings(project: Project) -> dict:
    from apps.comms.models import Meeting
    meetings = Meeting.objects.filter(project=project).order_by("-meeting_date")
    headers = ["Title", "Type", "Date", "Location", "Chaired By", "Actions Count"]
    rows = []
    for m in meetings:
        rows.append([
            m.title, m.get_meeting_type_display(), _fmt(m.meeting_date),
            m.location or "", m.chaired_by.get_full_name() if m.chaired_by else "",
            m.actions.count(),
        ])
    return {"title": f"Meetings - {project.name}", "headers": headers, "rows": rows}


def assemble_documents(project: Project) -> dict:
    from apps.documents.models import Document
    docs = Document.objects.filter(project=project).order_by("-created_at")
    headers = ["Name", "Category", "Current Version", "File Name", "Size (bytes)", "Last Updated"]
    rows = []
    for d in docs:
        rows.append([
            d.name, d.get_category_display(), d.current_version_number,
            d.latest_file_name, d.latest_file_size, _fmt(d.last_uploaded_at),
        ])
    return {"title": f"Document Register - {project.name}", "headers": headers, "rows": rows}


def assemble_progress(project: Project) -> dict:
    """High-level project progress summary combining schedule + cost."""
    from apps.scheduling.models import ProjectTask
    from apps.cost.models import BudgetLine, Expense
    from django.db.models import Avg, Sum

    tasks = ProjectTask.objects.filter(project=project, is_parent=False)
    task_count = tasks.count()
    completed = tasks.filter(status="completed").count()
    avg_progress = tasks.aggregate(avg=Avg("progress"))["avg"] or 0
    budget = BudgetLine.objects.filter(project=project).aggregate(t=Sum("budget_amount"))["t"] or 0
    actual = Expense.objects.filter(project=project).aggregate(t=Sum("amount"))["t"] or 0

    headers = ["Metric", "Value"]
    rows = [
        ["Project", project.name],
        ["Status", project.get_status_display()],
        ["Total Tasks", task_count],
        ["Completed Tasks", completed],
        ["Avg Progress", f"{avg_progress:.1f}%"],
        ["Total Budget", _fmt(budget)],
        ["Actual Cost", _fmt(actual)],
        ["Cost Variance", _fmt(budget - actual)],
    ]
    return {"title": f"Project Progress - {project.name}", "headers": headers, "rows": rows}


# Report key -> assembler mapping
REPORT_ASSEMBLERS = {
    "progress": assemble_progress,
    "financial": assemble_financial,
    "schedule": assemble_schedule,
    "procurement": assemble_procurement,
    "safety": assemble_safety,
    "quality": assemble_quality,
    "labour": assemble_labour,
    "risk": assemble_risk,
    "rfi": assemble_rfi,
    "changes": assemble_changes,
    "punch": assemble_punch,
    "documents": assemble_documents,
    "meetings": assemble_meetings,
}


# ---------------------------------------------------------------------------
# Export generators
# ---------------------------------------------------------------------------

def generate_csv(data: dict) -> tuple[bytes, str]:
    """Generate CSV bytes from assembled report data."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(data["headers"])
    for row in data["rows"]:
        writer.writerow(row)
    content = buf.getvalue().encode("utf-8-sig")
    return content, "text/csv"


def generate_xlsx(data: dict) -> tuple[bytes, str]:
    """Generate Excel bytes from assembled report data."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = data["title"][:31]  # Excel sheet name limit

    # Title row
    ws.append([data["title"]])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(data["headers"]))
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])  # blank row

    # Headers
    header_fill = PatternFill(start_color="1e293b", end_color="1e293b", fill_type="solid")
    header_font = Font(bold=True, color="e2e8f0")
    for col_idx, header in enumerate(data["headers"], 1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font

    # Data rows
    for row in data["rows"]:
        ws.append([str(v) if v is not None else "" for v in row])

    # Auto-width (skip merged cells)
    from openpyxl.cell.cell import MergedCell
    for col in ws.columns:
        cells = [c for c in col if not isinstance(c, MergedCell)]
        if not cells:
            continue
        max_len = max((len(str(cell.value or "")) for cell in cells), default=10)
        ws.column_dimensions[cells[0].column_letter].width = min(max_len + 2, 50)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def generate_pdf(data: dict) -> tuple[bytes, str]:
    """Generate PDF bytes from assembled report data."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    buf = io.BytesIO()
    page_size = landscape(A4) if len(data["headers"]) > 5 else A4
    doc = SimpleDocTemplate(buf, pagesize=page_size, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph(data["title"], styles["Title"]))
    elements.append(Spacer(1, 12))

    # Table
    table_data = [data["headers"]] + [[str(v) if v is not None else "" for v in row] for row in data["rows"]]
    if not table_data or len(table_data) < 2:
        elements.append(Paragraph("No data available.", styles["Normal"]))
    else:
        n_cols = len(data["headers"])
        col_width = (page_size[0] - 60) / n_cols
        t = Table(table_data, colWidths=[col_width] * n_cols)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#334155")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elements.append(t)

    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"Generated by BuildPro", styles["Normal"]))

    doc.build(elements)
    return buf.getvalue(), "application/pdf"


def generate_docx(data: dict) -> tuple[bytes, str]:
    """Generate Word DOCX bytes from assembled report data."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = DocxDocument()
    doc.add_heading(data["title"], level=1)
    doc.add_paragraph("")

    if data["rows"]:
        table = doc.add_table(rows=1 + len(data["rows"]), cols=len(data["headers"]))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # Headers
        for i, h in enumerate(data["headers"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data
        for r_idx, row in enumerate(data["rows"]):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val) if val is not None else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
    else:
        doc.add_paragraph("No data available.")

    doc.add_paragraph("")
    doc.add_paragraph("Generated by BuildPro")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# Format -> generator mapping
EXPORT_GENERATORS = {
    "csv": generate_csv,
    "xlsx": generate_xlsx,
    "pdf": generate_pdf,
    "docx": generate_docx,
}


# Need this for the financial assembler
from django.db import models
